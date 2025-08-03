from docx import Document

# 从服务器信息中生成一个简单的 Linux 运维报告，包含服务器名称、CPU 使用率、内存使用率等基本信息。
# 服务器清单
servers = [
    {"name": "server01", "cpu": "30%", "memory": "60%"},
    {"name": "server02", "cpu": "50%", "memory": "75%"},
    {"name": "server03", "cpu": "20%", "memory": "55%"}
]

# 创建word文档对象
doc = Document()
# 添加一级标题
doc.add_heading("Server Check Report", level=1)
# 添加段落信息

for server in servers:
    doc.add_heading(f"Server Name: {server['name']}", level=2)
    doc.add_paragraph(f"CPU usage: {server['cpu']}")
    doc.add_paragraph(f"Memory usage: {server['memory']}")
    doc.add_paragraph("-"*20)

# 保存文档
doc.save(r".\data-ai\report_batch.docx")