from docx import Document

# 提前有一个日志文件，设置变量
log_file = r".\data-ai\system.log"

# 创建word文档对象
doc = Document()

doc.add_heading("System log analyis report", level=1)
doc.add_paragraph("Below are the error messages in log files")
doc.add_paragraph("-"*30)

# 读取日志文件
with open(log_file, 'r') as f:
    for line in f:
        if "ERROR" in line or "WARNING" in line:
            doc.add_paragraph(line.strip())

doc.save(r".\data-ai\Error-analysis.docx")