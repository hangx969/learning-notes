from docx import Document

# 创建磁盘报告
disk_data = {
    "server01": [
        {"mount": "/", "size": "50G", "used": "20G", "free": "30G", "usage": "58%"},
        {"mount": "/home", "size": "100G", "used": "80G", "free": "20G", "usage": "80%"}
    ],
    "server02": [
        {"mount": "/", "size": "70G", "used": "30G", "free": "40G", "usage": "68%"},
        {"mount": "/var", "size": "200G", "used": "150G", "free": "50G", "usage": "78%"}
    ]
}

doc = Document()
doc.add_heading("Disk report", level=1)

# 循环生成每台服务器磁盘信息
for server, disks in disk_data.items():
    doc.add_heading(f"Server Name: {server}", level=2)
    # 创建表格存放磁盘数据，先只创建一个表头，后面的内容动态去创建
    table = doc.add_table(rows=1, cols=5)
    #设置表头. 会把第一行的每个单元格打包成元组
    hdr_cells = table.rows[0].cells
    # 给元组赋值，相当于给表头的的每个单元格赋值
    hdr_cells[0].text = "Mount Point"
    hdr_cells[1].text = "Total Size"
    hdr_cells[2].text = "Used"
    hdr_cells[3].text = "Free"
    hdr_cells[4].text = "Used Precent"

    for disk in disks:
        # .cells获取新增行的单元格
        row_cells = table.add_row().cells
        row_cells[0].text = disk['mount']
        row_cells[1].text = disk['size']
        row_cells[2].text = disk['used']
        row_cells[3].text = disk['free']
        row_cells[4].text = disk['usage']

doc.save(r".\data-ai\disk-usage-report.docx")