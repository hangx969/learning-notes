from docx import Document
import matplotlib.pyplot as plt
# 控制word中的图标大小
from docx.shared import Inches

plt.rcParams['font.family'] = ['Consolas']

# 性能数据
performance_data = {
    "CPU": [20, 40, 60, 80, 50, 30],
    "Memory": [30, 50, 70, 60, 40, 20],
    "Time": ["10:00", "10:05", "10:10", "10:15", "10:20", "10:25"]
}

# 绘制图表
plt.plot(performance_data['Time'], performance_data['CPU'], label="CPU Usage", marker='o')
plt.plot(performance_data['Time'], performance_data['Memory'], label="Memory Usage", marker='x')

plt.xlabel("Time")
plt.ylabel("Usage%")
plt.title("Performance Monitor")
# 添加图例
plt.legend()
# 保存图片
plt.savefig(r".\data-ai\performance_trend.png")

# 创建word文档
doc = Document()
doc.add_heading("Server performence trand", level=1)
doc.add_paragraph("Below are the cpu and memory usage trend:")
doc.add_picture(r".\data-ai\performance_trend.png", width=Inches(5))
doc.save(r".\data-ai\monitoring-trend.docx")